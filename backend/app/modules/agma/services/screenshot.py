from playwright.sync_api import sync_playwright
from playwright.async_api import async_playwright, TimeoutError as PlaywrightTimeoutError
from fastapi import HTTPException, status
import asyncio
from dotenv import load_dotenv
from PIL import Image
from io import BytesIO
import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()


class GetTicketServices:
    def __init__(self):
        self.FRONTEND = os.getenv("PLAYWRIGHTFRONTEND")

    def generate(self, id: str, path: str):
        with sync_playwright() as p:
            browser = p.chromium.launch()

            page = browser.new_page(
                device_scale_factor=2,
            )

            page.goto(f"{self.FRONTEND}{path}?id={id}",
                      wait_until="networkidle")

            ticket = page.locator(f"#agma-ticket")
            ticket.wait_for()
            screenshot = ticket.screenshot(omit_background=True)
            browser.close()
        return screenshot

    async def generate_ticket(self, id: str, path: str):
        return await asyncio.to_thread(self.generate, id, path)

    async def screenshot_tickets(self, id: str, path: str, token: str, refresh_token: str):
        """
        CONVERTS A BULK OF TICKETS INTO 1 WORD DOCUMENTS
        FOR EVERY PAGE IN THE WEBAPP ROUTE

        Args:
            id (str): the selector for the tickets 
            path (str): the current route of the webapp
            token (str): access_token sends from the frontend
            refresh_token (str): refresh_token sends from the frontend

        Returns:
            Image (list): a lists of screenshots images"""
        
        try: 
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                context = await browser.new_context(
                    device_scale_factor=1,
                    base_url=self.FRONTEND,
                )

                await context.add_cookies(
                    [{
                        "name": "access_token",
                        "value": str(token),
                        "url": self.FRONTEND,
                    },
                        {
                        "name": "refresh_token",
                        "value": str(refresh_token),
                        "url": self.FRONTEND,

                    }]
                )

                page = await context.new_page()
                await page.add_init_script(
                    """localStorage.setItem("LoginStatus", "true");"""
                )

                await page.goto(path, wait_until="networkidle")
                await page.wait_for_timeout(2000)
                await page.wait_for_selector(id)
                tickets = page.locator(id)

                images = []
                count = await tickets.count()

                for ticket in range(count):
                    screenshot = await tickets.nth(ticket).screenshot(omit_background=True)
                    images.append(Image.open(BytesIO(screenshot)))
                await browser.close()
                return images   
        except PlaywrightTimeoutError as e:
            raise HTTPException(
                status_code=status.HTTP_408_REQUEST_TIMEOUT, detail=str(e))
        

    
    async def convert_to_doc_bulk(self, start_page:int, end_page:int, path:str, selector:str, token: str, refresh_token: str):
        """
        CONVERTS A BULK OF TICKETS INTO 1 WORD DOCUMENTS
        FOR EVERY PAGE IN THE WEBAPP ROUTE

        Args:
            start_page (int): the startpage of the webapp route
            end_page (int):  the endpage of the webapp route
            path (str): the current route of the webapp
            selector (str): the id for the tickets 
            token (str): access_token sends from the frontend
            refresh_token (str): refresh_token sends from the frontend

        Returns:
            BytesIO or Bytes: a stream of bytes documents
        """
        bulk_images = []
        for page in range(start_page, end_page+1):
            print(page)
            path_with_page = f"{path}&page={page}"
            images = await self.screenshot_tickets(selector,path_with_page, token, refresh_token)
            bulk_images.extend(images)
        
        return await self.convert_to_doc(bulk_images)
    
    

    async def convert_to_doc(self, images: list):
        """_summary_

        Args:
            images (list): _description_

        Returns:
            _type_: _description_
        """
        doc = BytesIO()

        document = Document()

        section = document.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        
        for i in range(0, len(images), 4):
            page_images = images[i:i+4]
            
            table = document.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for index, image in  enumerate(page_images):
                row = index // 2
                col = index % 2
                
                
                ceil = table.cell(row, col)
                para = ceil.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                buffer = BytesIO()
                image.save(buffer, format="PNG")
                buffer.seek(0)
                            
                
                run = para.add_run()
                run.add_picture(buffer, width=Inches(3.5))
                
            if i + 4 < len(images):
                document.add_page_break()
        document.core_properties.title = "Agma Tickets"
        document.core_properties.author = "Richard Rojo Jr."
        document.save(doc)
        doc.seek(0)
        return doc